from docx import Document

# Load the document
doc = Document('/mnt/data/JosephFuCV20240429.docx')

# Function to clean and reformat the resume
def reformat_resume(doc):
    new_doc = Document()
    
    # Title
    new_doc.add_heading('Joseph Fu Chung Fat (傅宗發)', level=1)
    
    # Contact Information
    new_doc.add_heading('Contact Information', level=2)
    contact_info = [
        "Address: Flat 22B Block 2 Parkland Villas, No.1 Tuen On Lane, Tuen Mun, NT",
        "Contact No.: 67757087",
        "Email: josef.fu@gmail.com"
    ]
    for info in contact_info:
        new_doc.add_paragraph(info)
    
    # Professional Qualifications
    new_doc.add_heading('Professional Qualifications', level=2)
    qualifications = [
        ("Microsoft Certified System Engineer (MCSE)", 1998),
        ("Cisco Certified Network Associate (CCNA)", 2001),
        ("Certified Ethical Hacker (CEH)", 2004),
        ("Computer Hacking Forensic Investigator (CHFI)", 2004),
        ("Vmware Certified Professional 4 (VCP 4)", 2011),
        ("Citrix Certified Administrator (CCA)", 2013),
        ("Vmware Certified Professional 4 (VCP 5)", 2014),
        ("AWS Certified Solution Architect (Associate)", 2019)
    ]
    for qual, year in qualifications:
        new_doc.add_paragraph(f"{qual} - {year}")
    
    # Education Background
    new_doc.add_heading('Education Background', level=2)
    new_doc.add_paragraph("Bachelor of Science (B.Sc.) in Materials Science and Engineering")
    new_doc.add_paragraph("National Cheng Kung University, Taiwan")
    
    # Working Experience
    new_doc.add_heading('Working Experience', level=2)
    
    experiences = [
        {
            "company": "Invesco Hong Kong Limited",
            "duration": "Dec 2011 – Aug 2023",
            "position": "Senior Engineer",
            "responsibilities": [
                "Implement AWS Migration project and deploy AWS resources.",
                "Evaluate and provide solutions for AWS services.",
                "Develop Python scripts for housekeeping outdated AWS resources.",
                "Provide training and guidance to Citrix team on Packer code and Bitbucket CI/CD pipeline.",
                "Act as SME for Windows Server OS in Cloud Engineering team.",
                "Assist and provide L2 support for Cloud Operations team regarding Windows issues.",
                "Develop Packer and Terraform scripts for deploying AWS resources using Bitbucket CI/CD pipeline.",
                "Setup Ansible environment for software installation on Packer script for Windows Server images.",
                "Provide server support to all servers in APAC regions.",
                "Update Windows server build image and perform Windows OS hardening.",
                "Install Windows server hardware and setup various servers.",
                "Perform data migration and server monitoring.",
                "Handle Active Directory Regional Administration and provide SOP documentation.",
                "Support Data Center operations and perform data backup and server room setup.",
                "Participate in BCP drills and vendor management."
            ],
            "projects": [
                "Develop Automated Framework for AWS resource provisioning.",
                "POC project with Citrix team for AWS migration.",
                "Data Center Consolidation Project.",
                "VSphere upgrade project.",
                "Invesco India IT Integration project.",
                "Evaluation of Cloud services.",
                "Building IT infrastructure for China JV office.",
                "Office relocation and data center design.",
                "Server re-stack and BCP for Melbourne office.",
                "Secured printing project and print server migration.",
                "Hong Kong BCP Datacenter relocation planning."
            ]
        },
        {
            "company": "IT Channel (HK) Ltd.",
            "duration": "May 2007 – Nov 2011",
            "position": "Wintel System Administrator (Seconded to Deutsche Bank AG Hong Kong)",
            "responsibilities": [
                "Maintain infrastructure and application servers for North Asia Region.",
                "Plan and implement office relocation and disaster recovery projects.",
                "Implement server virtualization and print server migration.",
                "Setup and deploy various projects including HP Blade servers and Citrix XenDesktop.",
                "Perform Windows system administration and Active Directory regional administration."
            ]
        },
        {
            "company": "Kanbay (HK) Ltd.",
            "duration": "August 2001 – May 2007",
            "position": "Consultant (Seconded to HongKong International Terminal Limited)",
            "responsibilities": [
                "Maintain ERP application and Citrix servers.",
                "Implement IT projects including McAfee ePO deployment and server migrations.",
                "Manage regional servers and data consolidation projects.",
                "Ensure system stability and performance.",
                "Provide technical support and training."
            ]
        },
        {
            "company": "Inlooktech.com Limited",
            "duration": "March 2000 – March 2001",
            "position": "Network Engineer",
            "responsibilities": [
                "Provide network and helpdesk support.",
                "Deploy ERP software and provide technical support.",
                "Supervise office migration and design network infrastructure.",
                "Implement ISO 9001 project and recruit MIS staff."
            ]
        },
        {
            "company": "Pacific Rim Solutions Ltd.",
            "duration": "August 1999 – February 2000",
            "position": "Application Engineer",
            "responsibilities": [
                "Support software and hardware solutions.",
                "Perform system integration and network device configuration."
            ]
        }
    ]
    
    for exp in experiences:
        new_doc.add_heading(exp['company'], level=3)
        new_doc.add_paragraph(f"Position: {exp['position']} ({exp['duration']})")
        new_doc.add_heading('Responsibilities:', level=4)
        for responsibility in exp['responsibilities']:
            new_doc.add_paragraph(f"- {responsibility}")
        if 'projects' in exp:
            new_doc.add_heading('Projects:', level=4)
            for project in exp['projects']:
                new_doc.add_paragraph(f"- {project}")
    
    # Language Skills
    new_doc.add_heading('Language Skills', level=2)
    languages = [
        "Mother Tongue: Cantonese",
        "Written and spoken English: Proficient",
        "Written and spoken Mandarin: Fluent"
    ]
    for lang in languages:
        new_doc.add_paragraph(lang)
    
    # Skill Set
    new_doc.add_heading('Skill Set', level=2)
    skills = [
        "AWS Cloud",
        "Experience in Gitlab / Bitbucket CI / CD pipeline",
        "Packer / Terraform / PowerShell / Python scripting",
        "ESM security monitoring and suppression",
        "Microsoft Distributed File System (DFS) implementation",
        "Virtual Desktop Infrastructure (VDI) implementation",
        "VMWare ESX and VSphere setup and maintenance",
        "Citrix XenDesktop integration with VDI",
        "Microsoft Windows configuration and maintenance",
        "Windows System administration and OS tuning",
        "Active Directory Regional Administration"
    ]
    for skill in skills:
        new_doc.add_paragraph(f"- {skill}")
    
    # Availability
    new_doc.add_heading('Date Available:', level=2)
    new_doc.add_paragraph("Immediate Availability")
    
    return new_doc

# Reformat the resume
reformatted_doc = reformat_resume(doc)

# Save the new document
reformatted_doc.save('/Users/josephfu/Python/JosephFu_CV_2024_0528.docx')
